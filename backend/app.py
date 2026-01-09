import os
import sys
from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import logging

from dotenv import load_dotenv

load_dotenv()

# Ensure agents module is in path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from agents.student import StudentAgent

# Configure Logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
import base64
from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.utils import secure_filename
import time
import json
from dotenv import load_dotenv
from openai import OpenAI

load_dotenv()

app = Flask(__name__)
CORS(app)

# Use the main project uploads folder (one level up from backend)
UPLOAD_FOLDER = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'uploads')
if not os.path.exists(UPLOAD_FOLDER):
    UPLOAD_FOLDER = 'uploads'  # Fallback to local folder
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# In-memory storage for session data
# Map session_id -> StudentAgent instance
active_agents = {}
# Initialize OpenAI client
client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))

# In-memory storage for session data (mock persistence)
sessions = {}

# Image extensions
IMAGE_EXTENSIONS = {'jpg', 'jpeg', 'png', 'gif', 'webp', 'bmp'}

@app.route('/api/notes', methods=['POST'])
def get_notes():
    """Get the content of uploaded notes file"""
    data = request.json
    filename = data.get('filename')
    
    if not filename:
        return jsonify({'error': 'No filename provided'}), 400
    
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    
    # Determine file format
    file_ext = os.path.splitext(filename)[1].lower().lstrip('.')
    
    try:
        content = ''
        
        # Handle image files
        if file_ext in IMAGE_EXTENSIONS:
            with open(filepath, 'rb') as f:
                image_data = base64.b64encode(f.read()).decode('utf-8')
            mime_type = f'image/{file_ext}'
            if file_ext == 'jpg':
                mime_type = 'image/jpeg'
            content = f'data:{mime_type};base64,{image_data}'
            return jsonify({
                'filename': filename,
                'content': content,
                'format': 'image'
            }), 200
        
        elif file_ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        
        elif file_ext == 'csv':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        
        elif file_ext == 'docx':
            try:
                from docx import Document
                doc = Document(filepath)
                content = '\n'.join([para.text for para in doc.paragraphs])
            except ImportError:
                return jsonify({'error': 'python-docx not installed'}), 500
        
        elif file_ext == 'pdf':
            try:
                from PyPDF2 import PdfReader
                reader = PdfReader(filepath)
                text_parts = []
                for page in reader.pages:
                    text_parts.append(page.extract_text() or '')
                content = '\n'.join(text_parts)
            except ImportError:
                return jsonify({'error': 'PyPDF2 not installed'}), 500
        
        else:
            # Try to read as text for unknown formats
            try:
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
            except:
                return jsonify({'error': 'Unsupported file format'}), 400
        
        return jsonify({
            'filename': filename,
            'content': content,
            'format': file_ext or 'txt'
        }), 200
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/notes/analyze', methods=['POST'])
def analyze_notes():
    """Analyze notes and return highlighted content with summary"""
    data = request.json
    filename = data.get('filename')
    
    if not filename:
        return jsonify({'error': 'No filename provided'}), 400
    
    filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(filepath):
        return jsonify({'error': 'File not found'}), 404
    
    file_ext = os.path.splitext(filename)[1].lower().lstrip('.')
    
    # Get the text content
    try:
        content = ''
        if file_ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        elif file_ext == 'docx':
            from docx import Document
            doc = Document(filepath)
            content = '\n'.join([para.text for para in doc.paragraphs])
        elif file_ext == 'pdf':
             return jsonify({'error': 'PDF files are not supported. Please convert to DOCX or TXT.'}), 400

        else:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
    except Exception as e:
        return jsonify({'error': f'Failed to read file: {str(e)}'}), 500
    
    if not content.strip():
        return jsonify({'error': 'File is empty'}), 400
    
    # Call OpenAI to analyze and highlight
    try:
        # Truncate content if too long to avoid token limits
        max_content_length = 15000
        truncated_content = content[:max_content_length] if len(content) > max_content_length else content
        
        prompt = f"""Analyze the following study notes. Return a JSON object with:
1. "summary": A concise summary (3-5 sentences) of the key points.
2. "highlighted_html": The text with HTML spans for highlighting:
   - <span class="highlight-definition">text</span> for definitions
   - <span class="highlight-concept">text</span> for key concepts
   - <span class="highlight-important">text</span> for important notes
   - <span class="highlight-example">text</span> for examples

IMPORTANT: Return ONLY valid JSON, no markdown, no code blocks. Escape all quotes and special characters properly.

Notes to analyze:

{truncated_content}"""

        response = client.chat.completions.create(
            model="gpt-4.1",
            messages=[
                {"role": "system", "content": "You are a study assistant. Return ONLY valid JSON with 'summary' and 'highlighted_html' keys. No markdown code blocks. Properly escape all special characters in JSON strings."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=16000
        )
        
        result_text = response.choices[0].message.content.strip()
        
        # Remove markdown code blocks if present
        if result_text.startswith('```'):
            lines = result_text.split('\n')
            # Remove first line (```json) and last line (```)
            lines = [l for l in lines if not l.strip().startswith('```')]
            result_text = '\n'.join(lines)
        
        # Try to find JSON object in response
        start_idx = result_text.find('{')
        end_idx = result_text.rfind('}')
        
        if start_idx != -1 and end_idx != -1:
            result_text = result_text[start_idx:end_idx + 1]
        
        try:
            result = json.loads(result_text)
        except json.JSONDecodeError:
            # Fallback: try to fix common issues
            import re
            # Replace unescaped newlines in strings
            result_text = re.sub(r'(?<!\\)\n', '\\n', result_text)
            try:
                result = json.loads(result_text)
            except:
                # Last resort: return original content with basic highlighting
                return jsonify({
                    'filename': filename,
                    'summary': 'Summary generation failed - displaying original content.',
                    'highlighted_html': f'<p>{content}</p>',
                    'format': file_ext
                }), 200
        
        return jsonify({
            'filename': filename,
            'summary': result.get('summary', ''),
            'highlighted_html': result.get('highlighted_html', content),
            'format': file_ext
        }), 200
        
    except json.JSONDecodeError as e:
        return jsonify({'error': f'Failed to parse AI response: {str(e)}'}), 500
    except Exception as e:
        return jsonify({'error': f'AI analysis failed: {str(e)}'}), 500

@app.route('/api/notes/list', methods=['GET'])
def list_notes():
    """List all uploaded files in the uploads folder"""
    try:
        files = []
        for filename in os.listdir(app.config['UPLOAD_FOLDER']):
            # Skip PDF files
            if filename.lower().endswith('.pdf'):
                continue
                
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if os.path.isfile(filepath):
                files.append({
                    'filename': filename,
                    'format': os.path.splitext(filename)[1].lower().lstrip('.'),
                    'size': os.path.getsize(filepath)
                })
        return jsonify({'files': files}), 200
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    files = request.files.getlist('file')
    if not files or files[0].filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    saved_filenames = []
    filepaths = []
    
    for file in files:
        if file:
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            saved_filenames.append(filename)
            filepaths.append(filepath)
            
    # For Hackathon simplicity, we might just have one global "user" context or temporary sessions.
    # But let's assume we initialize an agent per session later.
    # For /api/topics, since it's stateless in the old mock, we might need to assume a "default" session 
    # OR we return a session_id right here. 
    # However, the frontend flow seems to be Upload -> Topics -> Start Session.
    # So we'll store these filepaths temporarily or create an agent here and return an ID?
    # The current frontend might expect just filenames. 
    # Let's create a temporary agent just for topics generation or manage it better.
    # Better yet: The frontend sends 'filenames' to /api/start_session.
    # So for /api/topics, the frontend expects us to know what files were uploaded?
    # Actually, the mock /api/topics endpoint just waited.
    # Let's implementation:
    # 1. Upload saves files.
    # 2. /api/topics needs to know WHICH files. 
    # IF the frontend doesn't send filenames to /api/topics, we might have to assume "all recent files" or similar.
    # Let's look at the integration guide again: "Backend reads uploaded file".
    # We will assume for this MVP that we process the files sent in the request or just the latest ones.
    # But wait, /api/topics is a POST. Does it send data?
    # The mock used `request.json` but commented out `filename = data.get('filename')`.
    # Let's assume the frontend sends the list of filenames it just uploaded.
    
    return jsonify({'message': 'Files uploaded successfully', 'filenames': saved_filenames}), 200

@app.route('/api/topics', methods=['POST'])
def get_topics():
    data = request.json
    # We expect filenames to be passed here, or we use all files in upload folder?
    # Let's try to get filenames from request, if not, scan upload folder (hacky but works for demo).
    filenames = data.get('filenames', [])
    
    if not filenames:
        # Fallback: list all in uploads
        filenames = os.listdir(app.config['UPLOAD_FOLDER'])
    
    file_paths = [os.path.join(app.config['UPLOAD_FOLDER'], f) for f in filenames if os.path.exists(os.path.join(app.config['UPLOAD_FOLDER'], f))]
    
    if not file_paths:
        return jsonify({'error': 'No files found'}), 400

    # Create a temporary agent to generate topics
    # Note: This is expensive to rebuild vectorstore just for topics, but safe.
    try:
        temp_agent = StudentAgent()
        temp_agent.process_files(file_paths)
        topics = temp_agent.generate_topics()
        return jsonify({'topics': topics}), 200
    except Exception as e:
        logger.error(f"Error in get_topics: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/start_session', methods=['POST'])
def start_session():
    data = request.json
    topic = data.get('topic')
    filenames = data.get('filenames', []) # List of filenames
    
    if not filenames:
         filenames = os.listdir(app.config['UPLOAD_FOLDER'])
         
    file_paths = [os.path.join(app.config['UPLOAD_FOLDER'], f) for f in filenames]
    
    session_id = f"session_{os.urandom(4).hex()}"
    
    agent = StudentAgent()
    try:
        agent.process_files(file_paths)
        agent.start_learning(topic)
        active_agents[session_id] = agent
        return jsonify({'session_id': session_id}), 200
    except Exception as e:
         logger.error(f"Error starting session: {e}")
         return jsonify({'error': str(e)}), 500

@app.route('/api/continue_session', methods=['POST'])
def continue_session():
    data = request.json
    session_id = data.get('session_id')
    
    if session_id not in active_agents:
        return jsonify({'error': 'Session not found'}), 404
        
    agent = active_agents[session_id]
    agent.target_questions += 3
    return jsonify({'message': 'Session extended'}), 200

@app.route('/api/chat', methods=['POST'])
def chat():
    data = request.json
    session_id = data.get('session_id')
    user_answer = data.get('answer') # Can be None if it's the start
    
    if session_id not in active_agents:
        return jsonify({'error': 'Session not found'}), 404
    
    agent = active_agents[session_id]
    
    try:
        response = agent.chat(user_answer)
        return jsonify(response), 200
    except Exception as e:
        logger.error(f"Error in chat: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/api/evaluate', methods=['POST'])
def evaluate():
    data = request.json
    session_id = data.get('session_id')
    
    if session_id not in active_agents:
        return jsonify({'error': 'Session not found'}), 404
    
    agent = active_agents[session_id]
    
    try:
        evaluation = agent.evaluate()
        return jsonify(evaluation), 200
    except Exception as e:
        logger.error(f"Error in evaluation: {e}")
        return jsonify({'error': str(e)}), 500

@app.route('/health', methods=['GET'])
def health():
    return jsonify({'status': 'ok'}), 200

if __name__ == '__main__':
    app.run(debug=True, port=5000)
