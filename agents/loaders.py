from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
try:
    import pytesseract
except ImportError:
    pytesseract = None
from PIL import Image
import nbformat
import os

def load_file(path: str) -> list[Document]:
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        print(f"INFO: Skipping PDF file {path} as it is not supported.")
        return []

    if ext in [".jpg", ".png", ".jpeg"]:
        if pytesseract is None:
            print(f"WARNING: pytesseract not installed. Skipping image {path}")
            return []
        try:
            text = pytesseract.image_to_string(Image.open(path))
            return [Document(page_content=text, metadata={"source": path})]
        except Exception as e:
            print(f"WARNING: Failed to process image {path} with OCR: {e}")
            return []

    if ext == ".ipynb":
        nb = nbformat.read(path, as_version=4)
        cells = []
        for cell in nb.cells:
            if cell.cell_type in ("markdown", "code"):
                cells.append(cell.source)
        return [Document(
            page_content="\n".join(cells),
            metadata={"source": path}
        )]

    if ext == ".txt":
        return TextLoader(path, encoding='utf-8').load()

    if ext == ".docx":
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(path)
            text = "\n".join([para.text for para in doc.paragraphs])
            return [Document(page_content=text, metadata={"source": path})]
        except ImportError:
            raise ValueError("python-docx not installed")

    if ext == ".csv":
        # Treating CSV as plain text for simplicity in RAG
        return TextLoader(path, encoding='utf-8').load()

    raise ValueError(f"Nieobsługiwany format: {ext}")
